# -*- coding: utf-8 -*-
"""
Saves generated content into Word documents instead of dumping markdown to
the terminal, since John reviews and copies posts from a docx, not a console.
"""

import os
from datetime import date
from docx import Document


def append_to_doc(doc_path: str, heading: str, body: str) -> None:
    """Append a dated heading and body to doc_path, creating the file (with
    a title matching its filename) if it does not exist yet."""
    if os.path.exists(doc_path):
        doc = Document(doc_path)
    else:
        doc = Document()
        title = os.path.splitext(os.path.basename(doc_path))[0]
        doc.add_heading(title, level=1)

    doc.add_heading(f"{heading} -- {date.today().isoformat()}", level=2)
    for line in body.split("\n"):
        doc.add_paragraph(line)

    try:
        doc.save(doc_path)
    except PermissionError:
        raise SystemExit(
            f"\n[DOC] Could not save '{doc_path}'. Close it if it is open "
            "in Word, then run again.\n"
        )
